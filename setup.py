"""
setup.py

One-time setup script for the Job Search Automation Tool.
Run this once before using main.py.

What it does:
  1. Parses your resume + LinkedIn export → profile/parsed_profile.json
  2. Initializes the Google Sheet with headers + formatting
  3. Creates a basic resume_template.docx if you don't have one
  4. Validates your .env configuration

Usage:
    python setup.py
    python setup.py --skip-sheets    # Skip Google Sheets init
    python setup.py --skip-docs      # Skip document template creation
"""

import argparse
import json
import logging
import os
import sys
from pathlib import Path

PROJECT_ROOT = Path(__file__).parent
sys.path.insert(0, str(PROJECT_ROOT / "src"))

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    datefmt="%H:%M:%S",
    handlers=[logging.StreamHandler(sys.stdout)],
)
logger = logging.getLogger(__name__)


def load_env() -> None:
    try:
        from dotenv import load_dotenv
        load_dotenv(PROJECT_ROOT / ".env")
    except ImportError:
        logger.warning("python-dotenv not installed — reading from system environment")


def load_config() -> dict:
    import yaml
    config_path = PROJECT_ROOT / "config.yaml"
    if not config_path.exists():
        raise FileNotFoundError(f"config.yaml not found at {config_path}")
    with open(config_path, encoding="utf-8") as f:
        return yaml.safe_load(f)


# ---------------------------------------------------------------------------
# Step 1: Parse profile
# ---------------------------------------------------------------------------
def parse_profile() -> bool:
    from profile_parser import build_profile, save_profile

    resume_pdf = PROJECT_ROOT / "profile" / "resume.pdf"
    resume_docx = PROJECT_ROOT / "profile" / "resume.docx"
    linkedin_dir = PROJECT_ROOT / "profile" / "linkedin_export"
    output_path = str(PROJECT_ROOT / "profile" / "parsed_profile.json")

    # Determine resume path
    resume_path = ""
    if resume_pdf.exists():
        resume_path = str(resume_pdf)
        logger.info(f"Found resume PDF: {resume_pdf}")
    elif resume_docx.exists():
        resume_path = str(resume_docx)
        logger.info(f"Found resume DOCX: {resume_docx}")
    else:
        logger.warning(
            "No resume found. Place your resume at:\n"
            f"  {resume_pdf}  (PDF preferred)\n"
            f"  {resume_docx}  (DOCX also supported)\n"
            "Continuing with LinkedIn data only..."
        )

    # Check LinkedIn export
    if not linkedin_dir.exists() or not any(linkedin_dir.iterdir()):
        logger.warning(
            f"LinkedIn export directory is empty: {linkedin_dir}\n"
            "To get your LinkedIn export:\n"
            "  1. LinkedIn → Settings → Data Privacy → Get a copy of your data\n"
            "  2. Select 'Fast file' with Profile, Positions, Skills, Education\n"
            "  3. Extract the ZIP contents into profile/linkedin_export/\n"
            "Continuing with resume data only..."
        )

    if not resume_path and not linkedin_dir.exists():
        logger.error("No resume or LinkedIn data found. Cannot build profile.")
        logger.error("Please add at least one data source before running setup.")
        return False

    logger.info("Parsing profile...")
    profile = build_profile(resume_path or "", str(linkedin_dir))

    # Show summary
    logger.info("")
    logger.info("Profile parsed successfully:")
    logger.info(f"  Name:       {profile.get('name', '(not found)')}")
    logger.info(f"  Email:      {profile.get('email', '(not found)')}")
    logger.info(f"  Skills:     {len(profile.get('skills', []))} skills")
    logger.info(f"  Experience: {len(profile.get('experience', []))} positions")
    logger.info(f"  Education:  {len(profile.get('education', []))} entries")
    logger.info(f"  Resume text: {len(profile.get('raw_resume_text', ''))} chars")

    save_profile(profile, output_path)
    logger.info(f"Saved to: {output_path}")

    # Ask user to verify
    logger.info("")
    logger.info("IMPORTANT: Review the parsed profile for accuracy:")
    logger.info(f"  Open: {output_path}")
    logger.info("If experience bullets are missing or wrong, you can manually edit parsed_profile.json.")

    return True


# ---------------------------------------------------------------------------
# Step 2: Initialize Google Sheets
# ---------------------------------------------------------------------------
def reformat_sheets() -> bool:
    """Re-apply all formatting to both tabs (run after schema or style changes)."""
    sheets_id = os.getenv("GOOGLE_SHEETS_ID")
    sa_path = os.getenv("GOOGLE_SERVICE_ACCOUNT_JSON")
    if not sheets_id or not sa_path:
        logger.warning("Sheets credentials not set — skipping reformat")
        return False
    try:
        from sheets_updater import SheetsUpdater
        updater = SheetsUpdater(sheets_id=sheets_id, service_account_path=sa_path)
        updater.connect()
        updater.reformat()
        logger.info("Sheets reformatted successfully")
        return True
    except Exception as e:
        logger.error(f"Reformat failed: {e}")
        return False


def init_google_sheets() -> bool:
    sheets_id = os.getenv("GOOGLE_SHEETS_ID")
    sa_path = os.getenv("GOOGLE_SERVICE_ACCOUNT_JSON")

    if not sheets_id:
        logger.warning("GOOGLE_SHEETS_ID not set in .env — skipping Sheets initialization")
        return False
    if not sa_path or not Path(sa_path).exists():
        logger.warning(f"GOOGLE_SERVICE_ACCOUNT_JSON not found at '{sa_path}' — skipping Sheets initialization")
        return False

    try:
        from sheets_updater import SheetsUpdater
        updater = SheetsUpdater(sheets_id=sheets_id, service_account_path=sa_path)
        updater.connect()
        logger.info("Google Sheets initialized successfully")
        return True
    except Exception as e:
        logger.error(f"Google Sheets initialization failed: {e}")
        logger.info("Common fixes:")
        logger.info("  1. Make sure you've shared the sheet with your service account email")
        logger.info("  2. Enable Google Sheets API and Google Drive API in Google Cloud Console")
        logger.info("  3. Verify the service account JSON path is correct")
        return False


# ---------------------------------------------------------------------------
# Step 3: Create document templates
# ---------------------------------------------------------------------------
def create_document_templates() -> bool:
    """
    Create basic .docx templates with docxtpl variable placeholders.
    If the user's resume.docx is found, use it as the basis for the resume template.
    """
    templates_dir = PROJECT_ROOT / "templates"
    templates_dir.mkdir(exist_ok=True)

    resume_template_path = templates_dir / "resume_template.docx"
    cover_template_path = templates_dir / "cover_letter_template.docx"

    # --- Resume Template ---
    if resume_template_path.exists():
        logger.info(f"Resume template already exists: {resume_template_path}")
    else:
        try:
            _create_resume_template(str(resume_template_path))
            logger.info(f"Created resume template: {resume_template_path}")
        except Exception as e:
            logger.error(f"Failed to create resume template: {e}")
            return False

    # --- Cover Letter Template ---
    if cover_template_path.exists():
        logger.info(f"Cover letter template already exists: {cover_template_path}")
    else:
        try:
            _create_cover_letter_template(str(cover_template_path))
            logger.info(f"Created cover letter template: {cover_template_path}")
        except Exception as e:
            logger.error(f"Failed to create cover letter template: {e}")
            return False

    return True


def _create_resume_template(output_path: str) -> None:
    """Create a basic resume template with docxtpl placeholders."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        raise ImportError("python-docx is required: pip install python-docx")

    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    # Name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("{{name}}")
    run.bold = True
    run.font.size = Pt(18)

    # Contact
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("{{email}}").font.size = Pt(10)

    doc.add_paragraph()

    # Summary
    _add_template_section_header(doc, "PROFESSIONAL SUMMARY")
    doc.add_paragraph("{{summary}}")

    # Skills
    _add_template_section_header(doc, "SKILLS")
    doc.add_paragraph("{{skills_text}}")

    # Experience — docxtpl Jinja2 loop
    _add_template_section_header(doc, "EXPERIENCE")
    # Note: docxtpl uses {%p for ... %} for paragraph-level loops
    p = doc.add_paragraph()
    p.add_run("{%p for exp in experience %}")

    p = doc.add_paragraph()
    run = p.add_run("{{exp.title}}")
    run.bold = True
    run.font.size = Pt(11)
    p.add_run("  |  {{exp.company}}").font.size = Pt(11)

    doc.add_paragraph("{{exp.started_on}} – {{exp.finished_on}}")

    p = doc.add_paragraph()
    p.add_run("{%p for bullet in exp.bullets %}")
    bp = doc.add_paragraph(style="List Bullet")
    bp.add_run("{{bullet}}")
    p = doc.add_paragraph()
    p.add_run("{%p endfor %}")

    p = doc.add_paragraph()
    p.add_run("{%p endfor %}")

    # Education
    _add_template_section_header(doc, "EDUCATION")
    p = doc.add_paragraph()
    p.add_run("{%p for edu in education %}")

    p = doc.add_paragraph()
    run = p.add_run("{{edu.degree}} in {{edu.field}}")
    run.bold = True
    run.font.size = Pt(11)
    p.add_run("  |  {{edu.school}}").font.size = Pt(11)

    doc.add_paragraph("{{edu.end_date}}")

    p = doc.add_paragraph()
    p.add_run("{%p endfor %}")

    doc.save(output_path)


def _add_template_section_header(doc, text: str) -> None:
    """Add a section header for the template."""
    try:
        from docx.shared import Pt, RGBColor
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        pass

    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    try:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    except Exception:
        pass

    try:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "1F497D")
        pBdr.append(bottom)
        pPr.append(pBdr)
    except Exception:
        pass


def _create_cover_letter_template(output_path: str) -> None:
    """Create a basic cover letter template."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("python-docx is required: pip install python-docx")

    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Header
    p = doc.add_paragraph()
    run = p.add_run("{{name}}")
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph("{{email}}").style.font.size = Pt(10)
    doc.add_paragraph()
    doc.add_paragraph("{{date}}").style.font.size = Pt(11)
    doc.add_paragraph()
    doc.add_paragraph("Dear {{company}} Hiring Team,").style.font.size = Pt(11)
    doc.add_paragraph()

    # Body — cover letter paragraphs split by double newline
    # We use a simple placeholder; document_builder will replace with actual paragraphs
    body_para = doc.add_paragraph("{{cover_letter_body}}")
    body_para.style.font.size = Pt(11)

    doc.add_paragraph()
    doc.add_paragraph("Sincerely,").style.font.size = Pt(11)
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("{{name}}")
    run.bold = True
    run.font.size = Pt(11)

    doc.save(output_path)


# ---------------------------------------------------------------------------
# Step 4: Validate .env
# ---------------------------------------------------------------------------
def validate_env() -> bool:
    """Check that all required environment variables are set."""
    required = {
        "ANTHROPIC_API_KEY": "Anthropic API key (get from console.anthropic.com)",
        "GOOGLE_SHEETS_ID": "Google Sheet ID (from sheet URL)",
        "GOOGLE_SERVICE_ACCOUNT_JSON": "Path to Google service account JSON file",
        "EMAIL_SENDER": "Gmail address to send from",
        "EMAIL_PASSWORD": "Gmail App Password (not regular password)",
        "EMAIL_RECIPIENT": "Email address to receive notifications",
    }

    all_set = True
    for key, description in required.items():
        val = os.getenv(key, "")
        if not val or val.startswith("your_"):
            logger.warning(f"  NOT SET: {key} — {description}")
            all_set = False
        else:
            masked = val[:4] + "..." + val[-4:] if len(val) > 8 else "***"
            logger.info(f"  OK:      {key} = {masked}")

    return all_set


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def main():
    parser = argparse.ArgumentParser(
        description="One-time setup for the Job Search Automation Tool"
    )
    parser.add_argument("--skip-sheets", action="store_true", help="Skip Google Sheets initialization")
    parser.add_argument("--skip-docs", action="store_true", help="Skip document template creation")
    parser.add_argument("--skip-profile", action="store_true", help="Skip profile parsing")
    parser.add_argument("--reformat", action="store_true", help="Re-apply sheet formatting only (no profile/template changes)")
    args = parser.parse_args()

    load_env()

    # Shortcut: just reformat the sheets
    if args.reformat:
        logger.info("Reformatting Google Sheets...")
        reformat_sheets()
        return

    logger.info("=" * 60)
    logger.info("Job Search Tool — Setup")
    logger.info("=" * 60)

    # Create required directories
    for d in ["profile", "profile/linkedin_export", "templates", "output", "data", "logs"]:
        (PROJECT_ROOT / d).mkdir(parents=True, exist_ok=True)

    # Ensure seen_jobs.json exists
    seen_path = PROJECT_ROOT / "data" / "seen_jobs.json"
    if not seen_path.exists():
        with open(seen_path, "w") as f:
            f.write("{}")
        logger.info("Created data/seen_jobs.json")

    all_good = True

    # Step 1: Parse profile
    if not args.skip_profile:
        logger.info("")
        logger.info("STEP 1: Parsing resume + LinkedIn export...")
        result = parse_profile()
        if not result:
            all_good = False
    else:
        logger.info("STEP 1: Skipped (--skip-profile)")

    # Step 2: Initialize Google Sheets
    if not args.skip_sheets:
        logger.info("")
        logger.info("STEP 2: Initializing Google Sheets...")
        result = init_google_sheets()
        if not result:
            all_good = False
    else:
        logger.info("STEP 2: Skipped (--skip-sheets)")

    # Step 3: Create document templates
    if not args.skip_docs:
        logger.info("")
        logger.info("STEP 3: Creating document templates...")
        result = create_document_templates()
        if not result:
            all_good = False
    else:
        logger.info("STEP 3: Skipped (--skip-docs)")

    # Step 4: Validate environment
    logger.info("")
    logger.info("STEP 4: Validating .env configuration...")
    env_valid = validate_env()
    if not env_valid:
        logger.warning("Some environment variables are not set. Fill them in your .env file before running main.py.")
        all_good = False

    logger.info("")
    logger.info("=" * 60)
    if all_good:
        logger.info("Setup complete! You're ready to run:")
        logger.info("  python main.py --dry-run    # Test run (no writes)")
        logger.info("  python main.py              # Full run")
    else:
        logger.info("Setup completed with warnings. Address the issues above, then run:")
        logger.info("  python setup.py             # Re-run setup after fixing issues")
        logger.info("  python main.py --dry-run    # Test when ready")
    logger.info("=" * 60)


if __name__ == "__main__":
    main()
